# main.py

from flask import Flask, jsonify, redirect, render_template, url_for, flash, request
from flask_dance.contrib.github import github
from flask_dance.contrib.google import google
from flask_login import logout_user, login_required
import os
from app.models import db, login_manager
from app.oauth import github_blueprint, google_blueprint
import logging
from werkzeug.utils import secure_filename
from flask import send_from_directory
from io import StringIO
from summarizer import Summarizer
import re
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from google.cloud import language_v1
from flask_caching import Cache
import math
from keybert import KeyBERT
from PyDictionary import PyDictionary
from nltk.corpus import wordnet
from docx import Document
from google.cloud import vision
import io
from google.cloud import speech

config = {
    "DEBUG": True,          # some Flask specific configs
    "CACHE_TYPE": "SimpleCache",  # Flask-Caching related configs
    "CACHE_DEFAULT_TIMEOUT": 3600
}


UPLOAD_FOLDER = './uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif','docx','flac','wav','mp3'}

app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///./users.db"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = "supersecretkey"
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "supersekrit")
app.config.from_mapping(config)
cache = Cache(app)

app.register_blueprint(google_blueprint, url_prefix="/login")
app.register_blueprint(github_blueprint, url_prefix="/login")

db.init_app(app)
login_manager.init_app(app)

with app.app_context():
    db.create_all()


@app.route("/ping")
def ping():
    return jsonify(ping="pong")


@app.route("/")
def homepage():
    return render_template("home.html")

@app.route("/login")
def login():
    return render_template("login.html")

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/github")
def login_github():
    if not github.authorized:
        return redirect(url_for("github.login"))
    res = github.get("/user")
    username = res.json()["login"]
    return f"You are @{username} on GitHub"


@app.route("/google")
def login_google():
    if not google.authorized:
        return redirect(url_for("google.login"))
    res = google.get("/oauth2/v1/userinfo")
    username = res.json()["email"]
    return f"You are @{username} on Google"


@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for("homepage"))

@app.errorhandler(500)
def server_error(e):
    logging.exception('An error occurred during a request.')
    return """
    An internal error occurred: <pre>{}</pre>
    See logs for full stacktrace.
    """.format(e), 500

@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_file():
    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        file = request.files['file']
        # If the user does not select a file, the browser submits an
        # empty file without a filename.
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            return redirect(url_for('download_file', name=filename))
    return '''
<!DOCTYPE html>
<html lang="en">

<head>
    <title>Upload | IntelliMed</title>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <meta http-equiv="X-UA-Compatible" content="ie=edge" />
   <link href="https://unpkg.com/material-components-web@latest/dist/material-components-web.min.css" rel="stylesheet">
  <script src="https://unpkg.com/material-components-web@latest/dist/material-components-web.min.js"></script>
<link rel="stylesheet" href="{{ url_for('static', filename='ustyle.css') }}">
<style>
*,
*:before,
*:after {
  box-sizing: border-box;
}

html {
  font-size: 16px;
  font-family: "Roboto", sans-serif;
}

body,
html {
  margin: 0;
  padding: 0;
  width:100%;
  height: 100%;
  background-color: #132D60;
}

.header-title-bar {
  height: 8.125rem;
  background-color: #c5c7c945;
  display: flex;
  justify-content: center;
  align-items: center;
  font-size: 2rem;
  font-weight: 700;
  color: #434343;
  text-align: center;
}

.upload {
  width: 90%;
  max-width: 1400px;
  display: flex;
  justify-content: center;
  align-items: center;
  flex-direction: column;
  margin: 2em auto;
  margin-bottom: 1em;
  padding: 2em;
  color: #434343;
  cursor: pointer;
  border: 0.1875em dashed #c5c7c945;
}
.upload svg {
  width: 15em;
  transition: all 0.4s ease;
}
.upload svg path {
  transition: all 0.4s ease;
}
.upload .cloud-img {
  padding-top: 3em;
}
.upload input {
  display: none;
}
.upload .img-title {
    color: #fff;
  font-size: 1.5rem;
  font-weight: 700;
  line-height: 1.5em;
  position: relative;
}
.upload .img-title span.clarify {
  display: block;
  color: #999999;
  font-size: 1rem;
  position: absolute;
  right: 0;
  bottom: 0;
  transform: translateX(101%);
}
.upload .blue {
  color: #3588b9;
  text-decoration: none;
  cursor: pointer;
}
.upload .img-subtitle {
  color: #999999;
}
.upload.highlight {
  border-color: salmon;
}
.upload.highlight svg {
  transform: scale(1.1);
}
.upload.highlight svg path {
  fill: #209957;
}

.d-none {
  display: none;
}

.btn {
  height: 40px;
  border-radius: 8px;
  padding: 0 16px;
  -webkit-transition: 0.1s ease-in-out;
  transition: 0.1s ease-in-out;
  font-weight: 500;
  cursor: pointer;
  color: #fff;
  -webkit-transition: 0.1s linear;
  transition: 0.1s linear;
  background-color: #209957;
  position: relative;
  border: none;
  left: 50%;
  transform: translateX(-50%);
  margin-bottom: 2em;
}

.btn:disabled {
  opacity: 0.6;
}

.success {
  color: #155724;
  background-color: #d4edda;
  border-color: #c3e6cb;
  position: relative;
  padding: 0.75rem 1.25rem;
  margin-bottom: 1rem;
  border: 1px solid transparent;
  border-radius: 0.25rem;
  width: 80%;
  margin: 1em auto;
}
h1{
    font-family: "Helvetica", sans-serif;
    font-size: 1.5rem;
    color: #fff;
  }
  .fried{
    font-family: "Poppins", sans-serif;
    color:yellow;
  }
  a {
    color:#fff;
    text-decoration: none;
    text-transform: uppercase;
    font-family: "Poppins", sans-serif;
  }
  h2{
    font-family: "Poppins", sans-serif;
    color: #fff;
    font-size: 4rem;
   
  }
  .topnav .search-container{
     text-align: center;  
     
  
}
  .topnav .search-container button {
   
    background: #ddd;
    font-size: 17px;
    border: none;
    cursor: pointer;
    border-radius:0.5rem;
    height:2.75rem;
    width:5rem;
  }
  .topnav .search-container button:hover {
    background: #ccc;
    size: 3rem;
  }
  input{
      width:40rem;
      height:3rem;
      text-align:center;
      border-radius: 0.5rem;
      font-size: 1.5rem;
  }
  .banana input{
    width:5rem;
      height:5rem;
    font-size: 3rem; 
margin-top:5rem;
  position: absolute;
  top: 100%;
  left: 50%;
  -ms-transform: translate(-50%, -50%);
  transform: translate(-50%, -50%);
  }
  .potato button{
  
      background-color: #28313b;
      color: #fff;
      border-radius: 1rem;
      width: 9rem;
      height: 3.5rem;
      display: inline-block;
      text-decoration: none;
      text-align: center;
      -ms-transform: translate(-50%, -50%);
  transform: translate(-50%, -50%);
  position: absolute;
 font-size: 1.25rem;
  left: 50%;


  }
  .boxid2{
    width: 80rem;
    height: 40rem;
    background-color: whitesmoke;
    border-radius: 2rem;
    margin-top: -5rem;
    margin-left: 5rem;
    
    box-shadow: 10px 10px #cc3300;
  }

  .card-grad--shot {
    color: transparent;
    background: linear-gradient(to right, hsl(294, 96%, 50%), hsl(201, 96%, 50%));
    -webkit-background-clip: text;
            background-clip: text;
  }


main {
  margin: 100px auto;
  position: relative;
  padding: 10px 5px;
  background: hsla(0,0%,100%,.3);
  font-size: 20px;
  font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, 'Open Sans', 'Helvetica Neue', sans-serif;
  line-height: 1.5;
  border-radius: 10px;
  width: 60%;
  overflow: hidden;
}

main::before {
  content: '';
  margin: -35px;
  position: absolute;
  top: 0;
  right: 0;
  bottom: 0;
  left: 0;
  filter: blur(20px);
  z-index: -1;
}


footer {
  padding-top: 20px;
  font-weight: bold;
}

cite {
  font-style: normal;
  font-size: 22px;
}

@import url(https://fonts.googleapis.com/css?family=Lora);

</style>
</head>

<body id="ok">  
    <h1 style="text-align:left;float:left;">Intelli<span class = "card-grad--shot">Med</span></h1>
<br><br><br>
<h2 style = "text-align: center;"><span class="card-grad--shot">Your reports</span>, made <span class="fried">easy</span> </h2>
<h1 style = "text-align: center;">Let's start by uploading your document</h1>
<h3 style = "text-align: center;color:#fff;">Supports pdf.,docx.,wav., or img.(beta)</h3>
  <main>
    <form method=post enctype=multipart/form-data>
    <input type=file name=file>
    <input type=submit value=Upload class="btn">
    </form>
  </main>

<script src="https://cdnjs.cloudflare.com/ajax/libs/three.js/r121/three.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/vanta@latest/dist/vanta.net.min.js"></script>
<script>
VANTA.NET({
  el: "#ok",
  mouseControls: true,
  touchControls: true,
  gyroControls: false,
  minHeight: 200.00,
  minWidth: 200.00,
  scale: 1.00,
  scaleMobile: 1.00
})
</script>
</body>

</html>
    '''

@app.route('/uploads/<name>')
@cache.cached(timeout=3600)
@login_required
def download_file(name):
  if name.split('.')[1]=="pdf":
    output_string = StringIO()
    with open('./uploads/'+name, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)

    body = str(output_string.getvalue())
    model = Summarizer()
    result = model(body, min_length=60)
    full = ''.join(result)
    res = re.sub(' +', ' ', full)
    spacekiller = str(res)
    summary = spacekiller.replace('\n', ' ').replace('\r', '')
    doc=summary

  elif name.split('.')[1]=="docx":
    document = Document('./uploads/'+name)
    for para in document.paragraphs:
        body= para.text

    model = Summarizer()
    result = model(body, min_length=60)
    summary = ''.join(result)
    doc=summary

  elif name.split('.')[1]=="txt":
    body=open("./uploads/"+name,'r')

    model = Summarizer()
    result = model(body, min_length=60)
    summary = ''.join(result)
    doc=summary

  elif name.split('.')[1]=="jpg" or name.split('.')[1]=="png" or name.split('.')[1]=="jpeg":
    """Detects text in the file."""

    client = vision.ImageAnnotatorClient()

    with io.open("./uploads/"+name, 'rb') as image_file:
        content = image_file.read()

    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    texts = response.text_annotations

    textdes=""
    for text in texts:
        textdes+=text.description

    a_string = textdes
    split_string = a_string.split("Close", 1)

    body = split_string[0]

    if response.error.message:
        raise Exception(
            '{}\nFor more info on error messages, check: '
            'https://cloud.google.com/apis/design/errors'.format(
                response.error.message))

    model = Summarizer()
    result = model(body, min_length=60)
    summary = ''.join(result)
    doc=summary

  elif name.split('.')[1]=="wav" or name.split('.')[1]=="mp3" or name.split('.')[1]=="flac":

    client = speech.SpeechClient()

    with io.open("./uploads/"+name, "rb") as audio_file:
        content = audio_file.read()

    audio = speech.RecognitionAudio(content=content)
    config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.ENCODING_UNSPECIFIED,
        sample_rate_hertz=16000,
        language_code="en-US",
    )

    response = client.recognize(config=config, audio=audio)
    body=""
    # Each result is for a consecutive portion of the audio. Iterate through
    # them to get the transcripts for the entire audio file.
    for result in response.results:
        # The first alternative is the most likely one for this portion.
        summary=str(result.alternatives[0].transcript)
    

    doc= summary

  kw_model = KeyBERT('distilbert-base-nli-mean-tokens')
  processedkeywords = kw_model.extract_keywords(doc)

  keywords=[]
  for keyword in processedkeywords:
      keywords.append(keyword[0])

  meaning=[]

  for i in keywords:
      syns = wordnet.synsets(i)
      meaning.append(syns[0].definition())
  
  result = {}
  for key in keywords:
      for value in meaning:
          result[key] = value
          meaning.remove(value)
          break

  return render_template('summary.html',summary=str(summary),condition=sample_analyze_sentiment(summary),confidence = sample_classify_text(summary),keywords=result)

#sentiment analysis (how severe the condition is)
def sample_analyze_sentiment(text_content):
    """
    Analyzing Sentiment in a String

    Args:
      text_content The text content to analyze
    """

    client = language_v1.LanguageServiceClient()

    # text_content = 'I am so happy and joyful.'

    # Available types: PLAIN_TEXT, HTML
    type_ = language_v1.Document.Type.PLAIN_TEXT

    # Optional. If not specified, the language is automatically detected.
    # For list of supported languages:
    # https://cloud.google.com/natural-language/docs/languages
    language = "en"
    document = {"content": text_content, "type_": type_, "language": language}

    # Available values: NONE, UTF8, UTF16, UTF32
    encoding_type = language_v1.EncodingType.UTF8

    response = client.analyze_sentiment(request = {'document': document, 'encoding_type': encoding_type})
    # Get overall sentiment of the input document


    listok = list(str(response.document_sentiment.score))
    if len(listok)==3:
      responsestring = listok[0]+listok[1]+listok[2]
    elif len(listok)==4:
      responsestring = listok[0]+listok[1]+listok[2]+listok[3]
    else:
      responsestring = listok[0]+listok[1]+listok[2]+listok[3]
    responsefloat = float(responsestring)
        
    # Get sentiment for all sentences in the document
    # Get the language of the text, which will be the same as
    # the language specified in the request or, if not specified,
    # the automatically-detected language.
    if responsefloat in [-0.1,-0.2,-0.3,-0.4]:
        return "Your condition is slightly serious. Take good care of yourself and obey your doctor's instructions. Visit the clinic as many times as possible."
    elif responsefloat==0.1 or response.document_sentiment.score==0.0:
        return "You seem to be just fine to do most tasks. However, they may be a risk of illness. Take good rest and visit your doctor regularly."
    elif responsefloat>0.1:
        return "Way to go, you seem perfectly healthy! But do remember to eat healthy and exercise often. Visit your doctor atleast once in six months."
    else:
        return "Your condition is extremely serious! Take good care of yourself and visit the doctor immediately."

def sample_classify_text(text_content):
    """
    Classifying Content in a String

    Args:
      text_content The text content to analyze. Must include at least 20 words.
    """

    client = language_v1.LanguageServiceClient()

    # text_content = 'That actor on TV makes movies in Hollywood and also stars in a variety of popular new TV shows.'

    # Available types: PLAIN_TEXT, HTML
    type_ = language_v1.Document.Type.PLAIN_TEXT

    # Optional. If not specified, the language is automatically detected.
    # For list of supported languages:
    # https://cloud.google.com/natural-language/docs/languages
    language = "en"
    document = {"content": text_content, "type_": type_, "language": language}

    response = client.classify_text(request = {'document': document})
    # Loop through classified categories returned from the API
    for category in response.categories:
        # Get the name of the category representing the document.
        # See the predefined taxonomy of categories:
        # https://cloud.google.com/natural-language/docs/categories
        confidence = "Your doctor is "+str(math.ceil(category.confidence*100))+" per-cent confident with his diagnosis."
        return confidence


def keywordsfunc(text_content):
    doc=text_content

    kw_model = KeyBERT('distilbert-base-nli-mean-tokens')
    processedkeywords = kw_model.extract_keywords(doc)

    keywords=[]
    for keyword in processedkeywords:
        keywords.append(keyword[0])

    meaning=[]

    for i in keywords:
        syns = wordnet.synsets(i)
        meaning.append(syns[0].definition())
    
    result = {}
    for key in keywords:
        for value in meaning:
            result[key] = value
            meaning.remove(value)
            break    

    return result




"""
@app.route('/summary')
@login_required
def summary():
    return render_template('summary.html')
"""




if __name__ == '__main__':
    # This is used when running locally. Gunicorn is used to run the
    # application on Google App Engine. See entrypoint in app.yaml.
    app.run(host='127.0.0.1', debug=True)
